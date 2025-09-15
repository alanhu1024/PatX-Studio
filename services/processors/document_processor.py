"""文档处理器 - 处理各种格式的文档"""
import os
import logging
from typing import Dict, List, Optional, Any
import pdfplumber
import docx
from pathlib import Path

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """文档处理器，支持PDF、Word、TXT等格式"""
    
    SUPPORTED_FORMATS = {'.pdf', '.docx', '.doc', '.txt', '.text'}
    
    def __init__(self):
        """初始化文档处理器"""
        self.processed_docs = []
        
    def process_file(self, file_path: str) -> Dict[str, Any]:
        """
        处理单个文件
        
        Args:
            file_path: 文件路径
            
        Returns:
            处理后的文档数据
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        file_ext = file_path.suffix.lower()
        
        if file_ext not in self.SUPPORTED_FORMATS:
            raise ValueError(f"不支持的文件格式: {file_ext}")
        
        try:
            # 根据文件类型选择处理方法
            if file_ext == '.pdf':
                content = self._extract_pdf(file_path)
            elif file_ext in ['.docx', '.doc']:
                content = self._extract_word(file_path)
            elif file_ext in ['.txt', '.text']:
                content = self._extract_text(file_path)
            else:
                content = ""
            
            doc_data = {
                'filename': file_path.name,
                'filepath': str(file_path),
                'content': content,
                'file_type': file_ext,
                'language': self._detect_language(content),
                'length': len(content)
            }
            
            self.processed_docs.append(doc_data)
            return doc_data
            
        except Exception as e:
            logger.error(f"处理文件失败 {file_path}: {str(e)}")
            raise
    
    def process_files(self, file_paths: List[str]) -> List[Dict[str, Any]]:
        """
        批量处理文件
        
        Args:
            file_paths: 文件路径列表
            
        Returns:
            处理后的文档数据列表
        """
        results = []
        for file_path in file_paths:
            try:
                doc_data = self.process_file(file_path)
                results.append(doc_data)
            except Exception as e:
                logger.error(f"处理文件失败 {file_path}: {str(e)}")
                results.append({
                    'filename': os.path.basename(file_path),
                    'filepath': file_path,
                    'content': '',
                    'error': str(e)
                })
        return results
    
    def _extract_pdf(self, file_path: Path) -> str:
        """
        提取PDF文本
        
        Args:
            file_path: PDF文件路径
            
        Returns:
            提取的文本内容
        """
        text = []
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text.append(page_text)
        except Exception as e:
            logger.error(f"PDF提取失败: {str(e)}")
            raise
        
        return '\n'.join(text)
    
    def _extract_word(self, file_path: Path) -> str:
        """
        提取Word文档文本
        
        Args:
            file_path: Word文件路径
            
        Returns:
            提取的文本内容
        """
        try:
            doc = docx.Document(file_path)
            text = []
            
            # 提取段落
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            
            # 提取表格
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        text.append(' | '.join(row_text))
            
            return '\n'.join(text)
            
        except Exception as e:
            logger.error(f"Word文档提取失败: {str(e)}")
            raise
    
    def _extract_text(self, file_path: Path) -> str:
        """
        提取纯文本文件内容
        
        Args:
            file_path: 文本文件路径
            
        Returns:
            文件内容
        """
        try:
            # 尝试不同的编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'gb18030', 'big5', 'latin1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
            
            # 如果所有编码都失败，使用错误处理
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
                
        except Exception as e:
            logger.error(f"文本文件提取失败: {str(e)}")
            raise
    
    def _detect_language(self, text: str) -> str:
        """
        检测文本语言
        
        Args:
            text: 文本内容
            
        Returns:
            语言代码 (zh, en, etc.)
        """
        # 简单的语言检测逻辑
        chinese_chars = sum(1 for char in text if '\u4e00' <= char <= '\u9fff')
        
        if chinese_chars / len(text) > 0.3:
            return 'zh'
        else:
            return 'en'
    
    def clear(self):
        """清空已处理的文档"""
        self.processed_docs = []
    
    def get_processed_docs(self) -> List[Dict[str, Any]]:
        """获取所有已处理的文档"""
        return self.processed_docs
    
    def search_in_docs(self, keyword: str) -> List[Dict[str, Any]]:
        """
        在已处理的文档中搜索关键词
        
        Args:
            keyword: 搜索关键词
            
        Returns:
            包含关键词的文档列表
        """
        results = []
        for doc in self.processed_docs:
            if keyword.lower() in doc.get('content', '').lower():
                # 找到关键词前后的上下文
                content = doc['content']
                keyword_pos = content.lower().find(keyword.lower())
                
                # 提取上下文（前后各100个字符）
                start = max(0, keyword_pos - 100)
                end = min(len(content), keyword_pos + len(keyword) + 100)
                context = content[start:end]
                
                results.append({
                    'filename': doc['filename'],
                    'context': context,
                    'position': keyword_pos
                })
        
        return results